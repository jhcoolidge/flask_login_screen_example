import pandas as pd
from google.cloud import bigquery
from google.oauth2 import service_account
from file_handlers import is_allowed_file, read_file
import docx
from docx.shared import RGBColor
import re

try:
    error_doc = docx.Document("error_logs.docx")
except:
    error_doc = docx.Document()
    error_doc.save("error_logs.docx")
    print("Previous file was corrupted or didn't exist - new file was created.")

# -----------------------------------------------------------------------------------------
# SETUP BIGQUERY
# -----------------------------------------------------------------------------------------

# Set up path to bigquery credential file and load the service account from it

# TODO: ENTER THE PATH TO YOUR KEY HERE
key_path = ''
credentials = service_account.Credentials.from_service_account_file(
    key_path, scopes=["https://www.googleapis.com/auth/cloud-platform"],
)

#TODO: ENTER YOUR DATASET NAME HERE
dataset_name = ''

# Set up BQ Client with given credentials
client = bigquery.Client(credentials=credentials, project=credentials.project_id)

# -----------------------------------------------------------------------------------------
# DATABASE FUNCTIONS
# -----------------------------------------------------------------------------------------


# Retrieve the user's password from BigQuery, if there is one
def get_user_pwd(username):
    query = "SELECT * FROM `{}.user_login` WHERE username = ? LIMIT 1;".format(dataset_name)

    # Load positional parameters
    job_config = bigquery.QueryJobConfig(
        query_parameters=[
            bigquery.ScalarQueryParameter(None, "STRING", username)
        ]
    )

    try:
        result = client.query(query, job_config=job_config)

        for row in result:
            return row["password"]
    except Exception as e:
        print("There was an issue retrieving user's password. Error: " + str(e))

    return None


# Store new user's information
def store_user_pwd(username, hashed_pwd):
    query = "INSERT INTO `{}.users.user_login` VALUES (?, ?);".format(dataset_name)

    # Load positional parameters
    job_config = bigquery.QueryJobConfig(
        query_parameters=[
            bigquery.ScalarQueryParameter(None, "STRING", username),
            bigquery.ScalarQueryParameter(None, "BYTES", hashed_pwd),
        ]
    )

    try:
        client.query(query, job_config=job_config)
    except Exception as e:
        print("There was a problem saving the user's password. " + str(e))
        return False
    return True


# Upload data from uploaded file
def upload_data_from_file(file):
    # Get the dataset object
    dataset = client.get_dataset('{}.uploaded_data'.format(dataset_name))

    # Create a reference to a new table with the filename as the table's name
    table = dataset.table('upload_data')

    file_extension = file.filename.split('.')[-1]

    if is_allowed_file(file_extension):
        try:
            file_dataframe = read_file(file)

            # Fire the load job and wait for the result
            load_job = client.load_table_from_dataframe(file_dataframe, table,
                                                        job_config=setup_job_config(file_extension))
            load_job.result()

            if load_job.errors:
                get_bad_rows_from_errors(load_job.errors, file_dataframe)
                message = "Upload was successful, but there were non-fatal errors. Check error logs for more details."
            else:
                message = "Upload was successful."
        except Exception as e:
            message = "Upload was not successful: " + str(e)
    else:
        message = "File is not currently able to be handled by the program. Accepted format: csv"

    return message


# Get bad rows and write them to a file
def get_bad_rows_from_errors(errors, file_dataframe):
    paragraph = error_doc.add_paragraph()
    paragraph.alignment = 0

    row_number_index = 0
    for error_index, error in enumerate(errors):

        # Find value which caused the error
        error_value = re.findall(r"'(.*?)'", error['message'])[0]

        # Find the values in the row containing the erroneous value and get the row number
        found_error_records = file_dataframe[file_dataframe.isin([error_value]).any(axis=1)].values
        row_numbers = file_dataframe[file_dataframe.isin([error_value]).any(axis=1)].index

        # Cycle trough all rows with the value that threw the error (if there are multiple)
        for record in found_error_records:

            # Format the string to be printed to the docx
            error_string = """
            ******************************************
            Error # - {} Error Message - {}
            Row # {}
            """.format(error_index, error['message'], row_numbers[row_number_index] + 1)
            row_number_index += 1

            # Write the log to the paragraph
            paragraph.add_run(error_string)

            # Cycle through each value in the record to apply a specific font color to each one.
            for value in record:
                run = paragraph.add_run("| {}, ".format(str(value)))
                if value == error_value:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                else:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black color

            # End the error log
            paragraph.add_run("\n******************************************")

        row_number_index = 0

    # Get all the rows containing null values in required columns and get their row numbers
    null_record_dataframe = file_dataframe[file_dataframe.isna().any(axis=1)].values
    null_record_row_numbers = file_dataframe[file_dataframe.isna().any(axis=1)].index

    null_row_number_index = 0

    # Cycle trough all records found having a null value
    for record in null_record_dataframe:

        # Format the error string
        error_string = """
                ******************************************
                Null values found in required columns
                Row # {}
                """.format(null_record_row_numbers[null_row_number_index] + 1)
        null_row_number_index += 1

        # Write the error string to the docx
        paragraph.add_run(error_string)

        # Cycle through the values in the record to apply color
        for value in record:
            run = paragraph.add_run("| {}, ".format(str(value)))
            if pd.isnull(value):
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
            else:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black color

        # End the error string
        paragraph.add_run("\n******************************************")

    # Save the final document
    error_doc.save("error_logs.docx")


# Setup BQ job configuration
def setup_job_config(file_extension):
    # Set up job configuration and schema
    job_config = bigquery.job.LoadJobConfig(
        schema=[
            bigquery.SchemaField("name", "STRING"),
            bigquery.SchemaField("salary", "INT64"),
            bigquery.SchemaField("emp_id", "INT64"),
            bigquery.SchemaField("job_title", "STRING"),
            bigquery.SchemaField("years_exp", "INT64"),
        ],
        max_bad_records=1000,

    )

    # Overwrite current table contents
    job_config.write_disposition = bigquery.WriteDisposition.WRITE_TRUNCATE

    if file_extension.lower() == "csv":
        job_config.source_format = bigquery.SourceFormat.CSV
    elif file_extension.lower() == "json":
        job_config.source_format = bigquery.SourceFormat.NEWLINE_DELIMITED_JSON

    return job_config
